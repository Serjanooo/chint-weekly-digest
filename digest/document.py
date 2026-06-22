from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BLUE = RGBColor(0x00, 0x66, 0xB3)
DARK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x66, 0x72, 0x85)


def _font(run, size: float, bold: bool = False, color: RGBColor = DARK, italic: bool = False) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def _hyperlink(paragraph, text: str, url: str, size: float = 11):
    relationship = paragraph.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(int(size * 2)))
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0066B3")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([fonts, font_size, color, underline])
    run.append(props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def build_docx(data: dict, output: Path, start: date, end: date) -> Path:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = section.bottom_margin = Cm(2.1)
    section.left_margin = section.right_margin = Cm(2.25)
    section.header_distance = section.footer_distance = Cm(1.25)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(header.add_run("CHINT RUSSIA  |  РЕДАКЦИОННЫЙ ДАЙДЖЕСТ"), 8.5, True, MUTED)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    _font(title.add_run("Новостной дайджест"), 26, True, DARK)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    _font(subtitle.add_run(f"Материалы за {start.strftime('%d.%m.%Y')}–{end.strftime('%d.%m.%Y')}"), 11, color=MUTED)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    _font(p.add_run("10 вариантов заголовка"), 16, True, BLUE)
    for index, option in enumerate(data["title_options"], 1):
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Cm(0.7)
        paragraph.paragraph_format.first_line_indent = Cm(-0.45)
        paragraph.paragraph_format.space_after = Pt(4)
        _font(paragraph.add_run(option), 11, index == 1)

    document.add_section(WD_SECTION.NEW_PAGE)
    heading = document.add_paragraph()
    heading.paragraph_format.space_after = Pt(6)
    _font(heading.add_run("Готовый пост"), 16, True, BLUE)
    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(14)
    _font(note.add_run("Основной вариант заголовка выделен первым. Перед публикацией проверьте ссылки."), 9.5, color=MUTED, italic=True)

    post_title = document.add_paragraph()
    post_title.paragraph_format.space_after = Pt(12)
    _font(post_title.add_run(data["title_options"][0]), 14, True, DARK)

    for story in data["stories"]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.keep_together = True
        _font(paragraph.add_run("🔵 "), 11, True, BLUE)
        before, link_text, after = story["text"].partition(story["link_text"])
        _font(paragraph.add_run(before), 11)
        _hyperlink(paragraph, link_text, story["url"])
        _font(paragraph.add_run(after), 11)

    hashtag = document.add_paragraph()
    hashtag.paragraph_format.space_before = Pt(4)
    _font(hashtag.add_run("#CHINT_Новости"), 11, True, BLUE)

    footer = document.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(footer.add_run("Сформировано локально с помощью Codex CLI"), 8, color=MUTED)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output
